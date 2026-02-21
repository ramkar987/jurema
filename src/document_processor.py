"""
Processador de documentos: extrai texto de PDF, DOCX e TXT.
"""

from pathlib import Path
from typing import Optional


class DocumentProcessor:
    """Carrega arquivos e retorna o conteúdo como texto puro."""

    TIPOS_SUPORTADOS = [".pdf", ".docx", ".txt"]

    def load_file(self, file_path: str) -> Optional[str]:
        """
        Carrega um arquivo e retorna seu conteúdo como string.

        Args:
            file_path: Caminho completo para o arquivo

        Returns:
            Texto extraído ou None se falhar
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix not in self.TIPOS_SUPORTADOS:
            return None

        try:
            if suffix == ".pdf":
                return self._load_pdf(str(path))
            elif suffix == ".docx":
                return self._load_docx(str(path))
            elif suffix == ".txt":
                return self._load_txt(str(path))
        except Exception as e:
            print(f"[DocumentProcessor] Erro ao carregar '{path.name}': {e}")
            return None

    def _load_pdf(self, file_path: str) -> str:
        """Extrai texto de PDF com pypdf, preservando número de página."""
        from pypdf import PdfReader

        reader = PdfReader(file_path)
        pages = []

        for num, page in enumerate(reader.pages, start=1):
            texto = page.extract_text()
            if texto and texto.strip():
                pages.append(f"[Página {num}]\n{texto.strip()}")

        if not pages:
            raise ValueError("PDF sem texto extraível (pode ser imagem escaneada).")

        return "\n\n".join(pages)

    def _load_docx(self, file_path: str) -> str:
        """Extrai texto de DOCX incluindo parágrafos e tabelas."""
        from docx import Document

        doc = Document(file_path)
        partes = []

        # Parágrafos normais
        for para in doc.paragraphs:
            if para.text.strip():
                partes.append(para.text.strip())

        # Conteúdo de tabelas
        for tabela in doc.tables:
            for linha in tabela.rows:
                celulas = [c.text.strip() for c in linha.cells if c.text.strip()]
                if celulas:
                    partes.append(" | ".join(celulas))

        return "\n\n".join(partes)

    def _load_txt(self, file_path: str) -> str:
        """Lê arquivo de texto, tentando UTF-8 e depois Latin-1."""
        for encoding in ("utf-8", "latin-1", "cp1252"):
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise ValueError(f"Não foi possível decodificar: {file_path}")
